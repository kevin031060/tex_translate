
import hashlib
import time
import requests
import os
import re
import sys
import docx
from docx.shared import RGBColor
import pickle




if __name__ == '__main__':
    # print(translate())
    filename = "file/SupplementalFile.tex"
    todocx = False

    if(re.search('.tex$',filename)==None):
        sys.exit('The input should be .tex file. Exit.')

    print('LaTeX file:',filename)

    with open(filename, 'r') as source_file:
        source = source_file.read()

    ### Search for possible token conflicts
    conflicts=re.findall('\[ *[012][\.\,][0-9]+\]',source)
    if(conflicts!=[]):
        print('Token conflicts detected: ',conflicts)
        sys.exit('Tokens may overlap with the content. Change tokens or remove the source of conflict.')
    else:
        print('No token conflicts detected. Proceeding.')

    #%% replace begin and end
    bdoc=re.search(r'\\begin{document}',source)
    edoc=re.search(r'\\end{document}',source)
    #%%
    x = {}
    i = 0
    txt = source

    x[f'[1.{i}]'] = txt[:bdoc.end()]
    i+=1
    x[f'[1.{i}]'] = txt[edoc.start():]
    i+=1
    #%%
    for key in x.keys():
        txt = txt.replace(x[key], key)

    def repl_comment(obj):
        global i_comments
        i_comments += 1
        return f'[1.{i_comments-1}]'

    #%%
    ### replace comments
    recomment = re.compile(r'(?<!\\)[%].*')
    i_comments = i
    for m in recomment.finditer(txt):
        x[f'[1.{i}]'] = m.group()
        i+=1
    txt = recomment.sub(repl_comment, txt)
    # comments.append(m.group())

    #%% replace begin{} environment
    start_values=[]
    end_values=[]
    for m in re.finditer(r'\\begin{ *algorithm\** *}|\\begin{ *equation\** *}|\\begin{ *figure\** *}|\\begin{ *eqnarray\** *}|\\begin{ *multline\** *}'
                         +r'|\\begin{ *thebibliography *}|\\begin{ *verbatim\** *}|\\begin{ *table\** *}|\\begin{ *subequations\** *}|\\begin{ *align\** *}'
                         +r'|\\begin{ *displaymath\** *}|\\begin{ *gather\** *}',txt):
        start_values.append(m.start())
    for m in re.finditer(r'\\end{ *algorithm\** *}|\\end{ *equation\** *}|\\end{ *figure\** *}|\\end{ *eqnarray\** *}|\\end{ *multline\** *}'
                         +r'|\\end{ *thebibliography *}|\\end{ *verbatim\** *}|\\end{ *table\** *}|\\end{ *subequations\** *}|\\end{ *align\** *}'
                         +r'|\\end{ *displaymath\** *}|\\end{ *gather\** *}',txt):
        end_values.append(m.end())
    nitems=len(start_values)
    assert(len(end_values)==nitems)

    #%%
    i_eq_start = i
    for neq in range(nitems):
        x[f'[1.{i}]'] = txt[start_values[neq]:end_values[neq]]
        i+=1
    i_eq_end = i
    for i_eq in range(i_eq_start, i_eq_end):
        txt = txt.replace(x[f'[1.{i_eq}]'], f'[1.{i_eq}]')
    #%% replace latex comm
    def repl_command(obj):
        global i_command
        i_command += 1
        return f'[1.{i_command-1}]'
    recommand = re.compile(r'___GTEXFIXCOMMENT[0-9]*___|\\title|\\chapter\**|\\section\**|\\subsection\**|\\subsubsection\**|~*\\footnote[0-9]*|(\$+)(?:(?!\1)[\s\S])*\1|~*\\\w*\s*{[^}]*}\s*{[^}]*}|~*\\\w*\s*{[^}]*}|~*\\\w*')
    i_command = i
    for m in recommand.finditer(txt):
        x[f'[1.{i}]'] = m.group()
        i+=1
    txt = recommand.sub(repl_command, txt)

    filename_base = filename.split('.')[0]
    with open(f'{filename_base}_latexitem', 'wb') as fp:
        pickle.dump(x, fp)
    for key in x.keys():
        print(key, x[key])
        source = source.replace(key, x[key])
    #%%
    if not todocx:
        ### Save the processed output to .txt file
        limit=30000 # Estimated Google Translate character limit
        filebase = re.sub('.tex$','',filename)
        start=0
        npart=0
        for m in re.finditer(r'\.\n',txt):
            if(m.end()-start<limit):
                end=m.end()
            else:
                output_filename = filebase+'_%d.txt'%npart
                npart+=1
                with open(output_filename, 'w') as txt_file:
                    txt_file.write(txt[start:end])
                print('Output file:',output_filename)
                start=end
                end=m.end()
        output_filename = filebase+'_%d.txt'%npart
        with open(output_filename, 'w') as txt_file:
            txt_file.write(txt[start:])
        print('Output file:',output_filename)
        print('Supply the output file(s) to Google Translate')
        with open(f'{filename_base}.txt', 'wb') as fp:
             pickle.dump(txt, fp)
    else:
        doc = docx.Document()
        rr=doc.add_paragraph()
        rr.add_run(txt)
        doc.save(f"{filename_base}.docx")
